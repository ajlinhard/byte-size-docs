"""
markdown_to_docx.py
Convert Markdown text to a formatted .docx file using python-docx.

Supported Markdown syntax:
  # H1  ## H2  ### H3  #### H4
  **bold**  *italic*  ***bold+italic***  `inline code`
  - / * / + unordered lists  (up to 2 levels)
  1. ordered lists
  > blockquote
  --- horizontal rule
  Plain paragraphs
"""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Inline formatting ────────────────────────────────────────────────────────

INLINE_RE = re.compile(
    r'(\*\*\*(?P<bolditalic>.+?)\*\*\*)'
    r'|(\*\*(?P<bold>.+?)\*\*)'
    r'|(\*(?P<italic>.+?)\*)'
    r'|(`(?P<code>.+?)`)'
)

def add_inline(paragraph, text):
    """Parse inline markdown and add styled runs to *paragraph*."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        # plain text before this match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        pos = m.end()

        if m.group('bolditalic'):
            run = paragraph.add_run(m.group('bolditalic'))
            run.bold = True
            run.italic = True
        elif m.group('bold'):
            run = paragraph.add_run(m.group('bold'))
            run.bold = True
        elif m.group('italic'):
            run = paragraph.add_run(m.group('italic'))
            run.italic = True
        elif m.group('code'):
            run = paragraph.add_run(m.group('code'))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

    # remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ── Horizontal rule ──────────────────────────────────────────────────────────

def add_horizontal_rule(doc):
    """Insert a paragraph with a bottom border to mimic ---."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── List helpers ─────────────────────────────────────────────────────────────

def add_list_item(doc, text, ordered=False, level=0):
    style = 'List Number' if ordered else 'List Bullet'
    if level > 0:
        style += f' {level + 1}'   # 'List Bullet 2', 'List Number 2', …
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        # Fallback if the numbered style isn't in the template
        p = doc.add_paragraph(style='List Bullet')
    add_inline(p, text)
    return p


# ── Document styles ──────────────────────────────────────────────────────────

def apply_document_styles(doc):
    """Set sensible defaults for Normal and Heading styles."""
    styles = doc.styles

    # Normal text: Arial 11 pt
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)

    heading_defs = [
        ('Heading 1', 20, True,  RGBColor(0x1F, 0x49, 0x7D)),
        ('Heading 2', 16, True,  RGBColor(0x2E, 0x74, 0xB5)),
        ('Heading 3', 13, True,  RGBColor(0x1F, 0x49, 0x7D)),
        ('Heading 4', 11, True,  RGBColor(0x2E, 0x74, 0xB5)),
    ]
    for name, size, bold, color in heading_defs:
        try:
            style = styles[name]
            style.font.name = 'Arial'
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = color
        except KeyError:
            pass


# ── Main converter ───────────────────────────────────────────────────────────

HEADING_RE  = re.compile(r'^(#{1,6})\s+(.*)')
UL_RE       = re.compile(r'^(\s*)[-*+]\s+(.*)')
OL_RE       = re.compile(r'^(\s*)\d+[.)]\s+(.*)')
BLOCKQUOTE  = re.compile(r'^>\s?(.*)')
HR_RE       = re.compile(r'^[-*_]{3,}\s*$')
TABLE_ROW_RE = re.compile(r'^\|.+\|$')
TABLE_SEP_RE = re.compile(r'^\|[-|: ]+\|$')

def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip('|').split('|')]


def add_table(doc, headers: list[str], rows: list[list[str]]):
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass

    # Header row — blue background, white bold text
    for j, text in enumerate(headers):
        cell = table.rows[0].cells[j]
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, cells in enumerate(rows):
        tr = table.rows[i + 1]
        for j in range(num_cols):
            text = cells[j] if j < len(cells) else ''
            add_inline(tr.cells[j].paragraphs[0], text)

    return table


def markdown_to_docx(markdown_text: str, output_path: str):
    doc = Document()
    apply_document_styles(doc)

    # Set US Letter page size with 1-inch margins
    section = doc.sections[0]
    section.page_width  = int(8.5 * 914400)
    section.page_height = int(11  * 914400)
    for attr in ('left_margin','right_margin','top_margin','bottom_margin'):
        setattr(section, attr, Inches(1))

    lines = markdown_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Blank line ──────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if HR_RE.match(line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ────────────────────────────────────────────────────────
        m = HEADING_RE.match(line)
        if m:
            level   = len(m.group(1))          # 1–6
            content = m.group(2).strip()
            style   = f'Heading {min(level, 4)}'
            p = doc.add_paragraph(style=style)
            add_inline(p, content)
            i += 1
            continue

        # ── Unordered list ──────────────────────────────────────────────────
        m = UL_RE.match(line)
        if m:
            indent  = len(m.group(1)) // 2
            content = m.group(2)
            add_list_item(doc, content, ordered=False, level=indent)
            i += 1
            continue

        # ── Ordered list ────────────────────────────────────────────────────
        m = OL_RE.match(line)
        if m:
            indent  = len(m.group(1)) // 2
            content = m.group(2)
            add_list_item(doc, content, ordered=True, level=indent)
            i += 1
            continue

        # ── Table ───────────────────────────────────────────────────────────
        if TABLE_ROW_RE.match(line):
            table_lines = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i]):
                table_lines.append(lines[i])
                i += 1
            # Expect: header row, separator row, then data rows
            if len(table_lines) >= 2:
                headers = _split_table_row(table_lines[0])
                data_rows = [
                    _split_table_row(ln)
                    for ln in table_lines[2:]
                    if not TABLE_SEP_RE.match(ln)
                ]
                add_table(doc, headers, data_rows)
            continue

        # ── Blockquote ──────────────────────────────────────────────────────
        m = BLOCKQUOTE.match(line)
        if m:
            try:
                p = doc.add_paragraph(style='Quote')
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
            add_inline(p, m.group(1))
            i += 1
            continue

        # ── Normal paragraph (collect continuation lines) ───────────────────
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not any([
            HEADING_RE.match(lines[i]),
            UL_RE.match(lines[i]),
            OL_RE.match(lines[i]),
            BLOCKQUOTE.match(lines[i]),
            HR_RE.match(lines[i]),
        ]):
            para_lines.append(lines[i])
            i += 1

        full_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        add_inline(p, full_text)

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── CLI entry-point ──────────────────────────────────────────────────────────

if __name__ == '__main__':
    SAMPLE_MD = """\
# Project Report

## Introduction

This document demonstrates **converting Markdown to Word** with *python-docx*.
It supports ***bold italic***, `inline code`, and other common elements.

## Features

- Headings H1–H4 with styled colors
- Unordered and ordered lists
  - Nested items at level 2
- **Bold**, *italic*, and ***combined*** inline styles
- `Inline code` with monospace font

## Steps to Reproduce

1. Install python-docx: `pip install python-docx`
2. Run this script: `python markdown_to_docx.py`
3. Open the generated **output.docx** in Word or LibreOffice

## Supported Elements

| Element       | Syntax                  | Notes                        |
|---------------|-------------------------|------------------------------|
| Heading       | `# H1` … `#### H4`     | Styled with blue colors      |
| **Bold**      | `**text**`              | Works inline                 |
| *Italic*      | `*text*`                | Works inline                 |
| `Code`        | `` `text` ``            | Monospace, red               |
| Unordered list| `- item`                | Up to 2 indent levels        |
| Ordered list  | `1. item`               | Up to 2 indent levels        |
| Blockquote    | `> text`                | Left-indented                |
| Table         | `\| col \| col \|`      | Header row styled blue/white |
| Horiz. rule   | `---`                   | Rendered as bottom border    |

## Blockquote Example

> *"Any fool can write code that a computer can understand.*
> *Good programmers write code that humans can understand."*
> — Martin Fowler

---

## Conclusion

The converter handles the **most common Markdown patterns** cleanly.
For complex documents (tables, images, footnotes), extend the parser or
use **pandoc** as a backend.
"""

    out = sys.argv[1] if len(sys.argv) > 1 else "/mnt/user-data/outputs/output.docx"
    markdown_to_docx(SAMPLE_MD, out)
